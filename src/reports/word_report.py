"""Word report generator — professional .docx benchmark report.

Uses python-docx for .docx generation. Creates:
- Cover page (title, date, author)
- Executive Summary (top performer, key findings)
- Per-model sections with metrics tables
- Conclusion with recommendations

Falls back to plain text if python-docx is not installed.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone

from src.core.logging import get_logger
from src.reports.excel_report import PortfolioReportData, ReportConfig

log = get_logger(__name__)


class WordReportGenerator:
    """Generate professional Word (.docx) reports."""

    def __init__(self, config: ReportConfig | None = None) -> None:
        self._config = config or ReportConfig()

    def generate(self, portfolios: list[PortfolioReportData]) -> bytes:
        """Generate Word document as bytes.

        Returns bytes that can be written to file or sent as HTTP response.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            log.warning("python_docx_not_installed", msg="Using fallback plain text output")
            return self._generate_text_fallback(portfolios)

        doc = Document()

        # Cover page
        self._write_cover_page(doc, Pt, WD_ALIGN_PARAGRAPH, RGBColor)

        # Executive summary
        self._write_executive_summary(doc, portfolios, Pt, WD_ALIGN_PARAGRAPH, RGBColor)

        # Per-model sections
        self._write_model_sections(doc, portfolios, Inches, Pt, RGBColor, WD_TABLE_ALIGNMENT)

        # Conclusion
        self._write_conclusion(doc, portfolios, Pt, WD_ALIGN_PARAGRAPH, RGBColor)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        log.info("word_report_generated", portfolios=len(portfolios))
        return buf.getvalue()

    def _write_cover_page(self, doc: object, Pt: type, WD_ALIGN_PARAGRAPH: object,
                          RGBColor: type) -> None:
        """Write the cover page with title, date, and author."""
        # Add spacing before title
        for _ in range(6):
            doc.add_paragraph("")

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(self._config.title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 78, 121)

        # Subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run("AI Trading Benchmark Analysis")
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(
            self._config.generated_at.strftime("%Y-%m-%d %H:%M UTC")
        )
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(128, 128, 128)

        # Author
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(f"Author: {self._config.author}")
        author_run.font.size = Pt(12)
        author_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_page_break()

    def _write_executive_summary(self, doc: object, portfolios: list[PortfolioReportData],
                                 Pt: type, WD_ALIGN_PARAGRAPH: object,
                                 RGBColor: type) -> None:
        """Write executive summary section."""
        heading = doc.add_heading("Executive Summary", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(31, 78, 121)

        if not portfolios:
            doc.add_paragraph("No portfolio data available for analysis.")
            return

        # Find top performers
        best_return = max(portfolios, key=lambda p: p.total_return_pct)
        best_sharpe = max(portfolios, key=lambda p: p.sharpe_ratio)
        best_caa = max(portfolios, key=lambda p: p.cost_adjusted_alpha)
        lowest_dd = min(portfolios, key=lambda p: p.max_drawdown)

        # Overview paragraph
        total_models = len({p.model for p in portfolios})
        total_markets = len({p.market for p in portfolios})
        doc.add_paragraph(
            f"This report summarizes the performance of {total_models} LLM model(s) "
            f"across {total_markets} market(s), comparing single-agent and multi-agent "
            f"architectures in virtual trading scenarios."
        )

        # Key findings
        doc.add_heading("Key Findings", level=2)

        findings = [
            f"Best Return: {best_return.model}/{best_return.architecture} "
            f"achieved {best_return.total_return_pct:+.2f}% total return "
            f"in the {best_return.market} market.",
            f"Best Risk-Adjusted: {best_sharpe.model}/{best_sharpe.architecture} "
            f"achieved the highest Sharpe ratio of {best_sharpe.sharpe_ratio:.4f}.",
            f"Best Cost-Adjusted Alpha: {best_caa.model}/{best_caa.architecture} "
            f"achieved CAA of {best_caa.cost_adjusted_alpha:.4f}, "
            f"indicating the best value after accounting for API costs.",
            f"Lowest Drawdown: {lowest_dd.model}/{lowest_dd.architecture} "
            f"had the smallest maximum drawdown at {lowest_dd.max_drawdown*100:.2f}%.",
        ]
        for finding in findings:
            doc.add_paragraph(finding, style="List Bullet")

        # Aggregate stats
        avg_return = sum(p.total_return_pct for p in portfolios) / len(portfolios)
        avg_sharpe = sum(p.sharpe_ratio for p in portfolios) / len(portfolios)
        total_cost = sum(p.total_api_cost for p in portfolios)

        doc.add_heading("Aggregate Statistics", level=2)
        doc.add_paragraph(f"Average Return: {avg_return:+.2f}%")
        doc.add_paragraph(f"Average Sharpe Ratio: {avg_sharpe:.4f}")
        doc.add_paragraph(f"Total API Cost: ${total_cost:.4f}")

    def _write_model_sections(self, doc: object, portfolios: list[PortfolioReportData],
                              Inches: type, Pt: type, RGBColor: type,
                              WD_TABLE_ALIGNMENT: object) -> None:
        """Write per-model detail sections with metrics tables."""
        doc.add_page_break()
        heading = doc.add_heading("Model Performance Details", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(31, 78, 121)

        # Group by model
        models: dict[str, list[PortfolioReportData]] = {}
        for p in portfolios:
            models.setdefault(p.model, []).append(p)

        for model_name, model_portfolios in models.items():
            doc.add_heading(f"Model: {model_name}", level=2)

            # Create metrics table
            table = doc.add_table(rows=1, cols=7)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            headers = ["Architecture", "Market", "Return %", "Sharpe",
                       "Max DD %", "Win Rate %", "Trades"]
            for idx, header_text in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = header_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)

            # Data rows
            for p in model_portfolios:
                row_cells = table.add_row().cells
                row_cells[0].text = p.architecture
                row_cells[1].text = p.market
                row_cells[2].text = f"{p.total_return_pct:+.2f}"
                row_cells[3].text = f"{p.sharpe_ratio:.4f}"
                row_cells[4].text = f"{p.max_drawdown * 100:.2f}"
                row_cells[5].text = f"{p.win_rate * 100:.1f}"
                row_cells[6].text = str(p.total_trades)

            doc.add_paragraph("")  # spacing

            # Model summary paragraph
            avg_return = sum(p.total_return_pct for p in model_portfolios) / len(model_portfolios)
            avg_sharpe = sum(p.sharpe_ratio for p in model_portfolios) / len(model_portfolios)
            model_cost = sum(p.total_api_cost for p in model_portfolios)

            doc.add_paragraph(
                f"{model_name} averaged {avg_return:+.2f}% return with a "
                f"mean Sharpe ratio of {avg_sharpe:.4f}. "
                f"Total API cost for this model: ${model_cost:.4f}."
            )

    def _write_conclusion(self, doc: object, portfolios: list[PortfolioReportData],
                          Pt: type, WD_ALIGN_PARAGRAPH: object, RGBColor: type) -> None:
        """Write conclusion and recommendations."""
        doc.add_page_break()
        heading = doc.add_heading("Conclusion & Recommendations", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(31, 78, 121)

        if not portfolios:
            doc.add_paragraph("No data available for conclusions.")
            return

        # Determine architecture comparison
        single_portfolios = [p for p in portfolios if p.architecture == "single"]
        multi_portfolios = [p for p in portfolios if p.architecture == "multi"]

        if single_portfolios and multi_portfolios:
            avg_single = sum(p.total_return_pct for p in single_portfolios) / len(single_portfolios)
            avg_multi = sum(p.total_return_pct for p in multi_portfolios) / len(multi_portfolios)
            cost_single = sum(p.total_api_cost for p in single_portfolios) / len(single_portfolios)
            cost_multi = sum(p.total_api_cost for p in multi_portfolios) / len(multi_portfolios)

            doc.add_heading("Architecture Comparison", level=2)
            doc.add_paragraph(
                f"Single-agent architectures averaged {avg_single:+.2f}% return "
                f"with ${cost_single:.4f} average API cost per agent. "
                f"Multi-agent architectures averaged {avg_multi:+.2f}% return "
                f"with ${cost_multi:.4f} average API cost per agent."
            )

            if avg_multi > avg_single:
                doc.add_paragraph(
                    "The multi-agent architecture showed a performance advantage, "
                    "though at higher API cost. Consider the cost-adjusted alpha (CAA) "
                    "metric when evaluating whether the additional cost is justified."
                )
            else:
                doc.add_paragraph(
                    "The single-agent architecture performed comparably or better, "
                    "suggesting that additional multi-agent complexity may not be "
                    "justified for the tested configurations."
                )

        # Recommendations
        doc.add_heading("Recommendations", level=2)
        best_caa = max(portfolios, key=lambda p: p.cost_adjusted_alpha)
        recommendations = [
            f"Deploy {best_caa.model}/{best_caa.architecture} for production use "
            f"based on superior cost-adjusted performance.",
            "Continue monitoring drawdown limits and implement automated circuit breakers.",
            "Consider expanding the benchmark to additional market conditions and time periods.",
            "Evaluate ensemble approaches combining top-performing model/architecture pairs.",
        ]
        for rec in recommendations:
            doc.add_paragraph(rec, style="List Number")

        # Footer
        doc.add_paragraph("")
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"Report generated by {self._config.author} on "
            f"{self._config.generated_at.strftime('%Y-%m-%d %H:%M UTC')}"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(160, 160, 160)

    def _generate_text_fallback(self, portfolios: list[PortfolioReportData]) -> bytes:
        """Plain text fallback when python-docx is not installed."""
        lines: list[str] = []
        lines.append(f"{'=' * 60}")
        lines.append(f"  {self._config.title}")
        lines.append(f"  Generated: {self._config.generated_at.strftime('%Y-%m-%d %H:%M UTC')}")
        lines.append(f"  Author: {self._config.author}")
        lines.append(f"{'=' * 60}")
        lines.append("")

        if not portfolios:
            lines.append("No portfolio data available.")
            return "\n".join(lines).encode("utf-8")

        # Executive summary
        lines.append("EXECUTIVE SUMMARY")
        lines.append("-" * 40)
        best_return = max(portfolios, key=lambda p: p.total_return_pct)
        lines.append(
            f"Best Return: {best_return.model}/{best_return.architecture} "
            f"({best_return.total_return_pct:+.2f}%)"
        )
        best_sharpe = max(portfolios, key=lambda p: p.sharpe_ratio)
        lines.append(
            f"Best Sharpe: {best_sharpe.model}/{best_sharpe.architecture} "
            f"({best_sharpe.sharpe_ratio:.4f})"
        )
        lines.append("")

        # Per-model details
        lines.append("MODEL DETAILS")
        lines.append("-" * 40)
        header = f"{'Model':<12} {'Arch':<8} {'Market':<8} {'Return%':>8} {'Sharpe':>8} {'MaxDD%':>8}"
        lines.append(header)
        lines.append("-" * len(header))
        for p in portfolios:
            lines.append(
                f"{p.model:<12} {p.architecture:<8} {p.market:<8} "
                f"{p.total_return_pct:>8.2f} {p.sharpe_ratio:>8.4f} "
                f"{p.max_drawdown * 100:>8.2f}"
            )
        lines.append("")

        # Conclusion
        lines.append("CONCLUSION")
        lines.append("-" * 40)
        best_caa = max(portfolios, key=lambda p: p.cost_adjusted_alpha)
        lines.append(
            f"Recommended: {best_caa.model}/{best_caa.architecture} "
            f"(CAA: {best_caa.cost_adjusted_alpha:.4f})"
        )

        return "\n".join(lines).encode("utf-8")
